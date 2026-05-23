from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from pathlib import Path


OUT_DIR = Path("/Users/baihe/Documents/bohack/outputs")
OUT_PATH = OUT_DIR / "此地有回声_AI城市时空棋局_京津双城首发版游戏文档.docx"


FONT_CN = "PingFang SC"
FONT_LATIN = "Calibri"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "666666"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
PALE_YELLOW = "FFF8E1"
PALE_GREEN = "EAF5EE"
PALE_RED = "FCECEC"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), "120")
    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(widths[idx]))
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run_font(run, size=None, bold=None, color=None):
    run.font.name = FONT_LATIN
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_para_format(p, before=0, after=6, line=1.25, keep_with_next=False):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    pf.keep_with_next = keep_with_next


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), FONT_LATIN)
    fonts.set(qn("w:hAnsi"), FONT_LATIN)
    fonts.set(qn("w:eastAsia"), FONT_CN)
    r_pr.append(fonts)
    new_run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def style_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.42)
    section.footer_distance = Inches(0.42)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT_LATIN
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[name]
        style.font.name = FONT_LATIN
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.18
        style.paragraph_format.keep_with_next = True

    footer_p = section.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_p.paragraph_format.space_before = Pt(0)
    footer_p.paragraph_format.space_after = Pt(0)
    run = footer_p.add_run("《此地有回声：AI 城市时空棋局》京津双城首发版游戏文档")
    set_run_font(run, 8.5, False, MUTED)


def add_title(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_format(p, before=20, after=8, line=1.08)
    r = p.add_run("《此地有回声》")
    set_run_font(r, 26, True, INK)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_format(p2, after=10, line=1.08)
    r = p2.add_run("AI 城市时空棋局")
    set_run_font(r, 20, True, BLUE)
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_format(p3, after=14, line=1.2)
    r = p3.add_run("京津双城首发版游戏设计文档")
    set_run_font(r, 14, False, DARK_BLUE)

    meta = [
        ("版本", "V0.1 / MVP 执行稿"),
        ("首发城市", "北京《中轴入局》；天津《海河来信》"),
        ("核心体验", "固定主线步行路线 + 时空骰随机事件 + 现场拍照触发 AI 角色剧场 + 文化卡牌 + 个人故事生成"),
        ("目标用途", "产品立项、比赛路演、原型设计、内容制作、技术开发对齐"),
    ]
    add_kv_table(doc, meta, widths=[1600, 7520], fill=LIGHT_BLUE)
    add_callout(
        doc,
        "一句话定位",
        "一款把北京和天津变成可行走剧情棋盘的 AI 文旅游戏 App。用户沿真实路线行走，通过定位、拍照和对话触发历史角色剧场，收集文化卡牌，最终生成自己的京津城市故事。",
        fill=PALE_GREEN,
    )


def add_callout(doc, label, body, fill=LIGHT_GRAY, spacer_after=True):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_width(table, [9120])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    set_para_format(p, after=3, line=1.2)
    r = p.add_run(label)
    set_run_font(r, 10.5, True, DARK_BLUE)
    p2 = cell.add_paragraph()
    set_para_format(p2, after=2, line=1.25)
    r = p2.add_run(body)
    set_run_font(r, 10.5, False, INK)
    if spacer_after:
        spacer = doc.add_paragraph()
        set_para_format(spacer, after=3)


def add_kv_table(doc, rows, widths=[1800, 7320], fill=LIGHT_GRAY):
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for k, v in rows:
        cells = table.add_row().cells
        cells[0].text = str(k)
        cells[1].text = str(v)
        set_cell_shading(cells[0], fill)
        for idx, cell in enumerate(cells):
            for p in cell.paragraphs:
                set_para_format(p, after=0, line=1.18)
                for run in p.runs:
                    set_run_font(run, 9.5, idx == 0, DARK_BLUE if idx == 0 else INK)
    set_table_width(table, widths)
    return table


def add_data_table(doc, headers, rows, widths, header_fill=LIGHT_BLUE, font_size=9.0):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        set_cell_shading(cell, header_fill)
        for p in cell.paragraphs:
            set_para_format(p, after=0, line=1.15)
            for run in p.runs:
                set_run_font(run, font_size, True, DARK_BLUE)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
            for p in cells[i].paragraphs:
                set_para_format(p, after=0, line=1.15)
                for run in p.runs:
                    set_run_font(run, font_size, False, INK)
    set_table_width(table, widths)
    return table


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        set_para_format(p, after=4, line=1.2)
        r = p.add_run(item)
        set_run_font(r, 10.5, False, INK)


def add_numbers(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        set_para_format(p, after=4, line=1.2)
        r = p.add_run(item)
        set_run_font(r, 10.5, False, INK)


def add_section_intro(doc, title, body):
    doc.add_heading(title, level=1)
    p = doc.add_paragraph()
    set_para_format(p, after=8, line=1.25)
    r = p.add_run(body)
    set_run_font(r, 10.5, False, INK)


def add_city_node(doc, node):
    doc.add_heading(node["title"], level=3)
    rows = [
        ("地点范围", node["place"]),
        ("用户输入", node["input"]),
        ("AI 判断", node["ai"]),
        ("剧场角色", node["role"]),
        ("开场台词", node["line"]),
        ("现实任务", node["task"]),
        ("获得卡牌", node["card"]),
        ("失败兜底", node["fallback"]),
    ]
    add_kv_table(doc, rows, widths=[1600, 7520], fill=LIGHT_GRAY)


def build_document():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style_document(doc)
    add_title(doc)

    doc.add_page_break()
    add_section_intro(
        doc,
        "1. 项目总览",
        "本项目不是传统景点导览，也不是随机走格子的城市版桌游。它的首发版本被收窄为北京、天津两座城市的实景剧情游戏：路线由产品预设，剧情由现场照片、定位、天气、时间和用户选择共同触发，AI 在史料护栏内扮演城市记忆中的人物，让游客成为故事参与者。",
    )
    add_kv_table(
        doc,
        [
            ("项目名称", "《此地有回声：AI 城市时空棋局》"),
            ("首发包名", "京津双城时空季"),
            ("北京篇", "《中轴入局》：从前门、故宫角楼、景山到鼓楼/什刹海，体验都城秩序与旧城生活。"),
            ("天津篇", "《海河来信》：从天津站、解放桥、意式风情区、海河岸到古文化街，体验开埠、建筑、商业与民俗。"),
            ("核心差异", "不是 AI 导游讲解，而是城市实景角色扮演；不是照片识别打卡，而是照片增强剧情；不是随机乱走，而是安全路线上的随机事件。"),
        ],
        widths=[1600, 7520],
        fill=LIGHT_BLUE,
    )
    add_callout(
        doc,
        "设计判断",
        "第一版只做两个城市、两条主线、每城五个关键棋格。每个棋格只打透一个 3-5 分钟的现场互动闭环：到达、拍照、剧场、任务、卡牌、记忆碎片。",
        fill=PALE_YELLOW,
    )

    doc.add_heading("1.1 目标用户", level=2)
    add_data_table(
        doc,
        ["用户类型", "核心动机", "产品要满足的体验"],
        [
            ("年轻游客", "想要比普通攻略更有参与感的城市漫游", "路线不累、内容有梗、照片能生成可分享成果。"),
            ("亲子家庭", "希望孩子在旅行中主动观察和提问", "任务清楚、对话安全、知识被故事包裹。"),
            ("研学/社团", "需要可组织、可复盘、可展示的城市文化任务", "多人协作、卡牌成果、路线和内容可配置。"),
            ("本地周末用户", "重新发现熟悉城市", "支线和隐藏事件能带来新鲜感。"),
        ],
        widths=[1600, 2520, 5000],
    )

    doc.add_heading("1.2 成功标准", level=2)
    add_bullets(
        doc,
        [
            "用户能在 90-150 分钟内完成一条城市主线，不需要提前理解复杂规则。",
            "每个地点都能从“现实观察”自然进入“AI 剧场”，而不是跳到百科讲解。",
            "用户愿意至少完成 3 次拍照、3 次对话选择、1 次文字或语音留言。",
            "最终生成的游记和海报足够个人化，能够看出用户走过哪里、选择了什么、留下了什么。",
            "内容团队可以用统一模板扩展新城市，而不是为每个点重新设计系统。",
        ]
    )

    add_section_intro(
        doc,
        "2. 核心玩法",
        "游戏采用“固定主线步行路线 + 时空骰随机事件 + 可选支线格子”的结构。路线由产品方预设，保证安全、顺路、可完成；随机性只影响每一站的剧情角度、任务形式和角色追问，不强迫用户在真实城市里乱走。",
    )
    doc.add_heading("2.1 单次游玩流程", level=2)
    add_numbers(
        doc,
        [
            "用户进入 App，选择城市：北京或天津。",
            "用户输入游玩时间、同行类型、兴趣主题和体力偏好。",
            "系统生成当日可走的“城市时空棋盘”，显示主线站点、支线站点、预计用时和安全提示。",
            "用户到达第一个棋格，GPS/围栏确认进入地点范围。",
            "用户拍摄现场照片，系统识别建筑元素、天气、角度和氛围。",
            "用户掷出“时空骰”，触发本地点的剧情事件类型。",
            "AI 角色登场，与用户进行 2-4 轮对话。",
            "用户完成现实任务，获得文化卡牌。",
            "系统记录路线、照片、选择、卡牌和留言。",
            "终局生成个人城市故事、卡牌册和分享海报。",
        ]
    )

    doc.add_heading("2.2 时空骰规则", level=2)
    add_data_table(
        doc,
        ["骰面", "触发内容", "体验效果"],
        [
            ("时辰", "根据早晚、光线、天气改变剧情氛围", "同一地点在晴天、夜晚、雨后会出现不同开场。"),
            ("风物", "把照片里的建筑元素变成剧情道具", "屋檐、拱窗、桥、河面、牌楼都会进入角色台词。"),
            ("来信", "历史角色向用户提出一个问题", "用户从游客变成被城市询问的人。"),
            ("市声", "触发声音观察任务", "录制钟声、人声、水声、街声，作为游记素材。"),
            ("转折", "出现历史处境选择", "用户选择影响获得的卡牌和最终故事标题。"),
            ("回声", "用户给这座城留下一句话", "终局生成“来自今天的回声”。"),
        ],
        widths=[1100, 3500, 4520],
    )

    doc.add_heading("2.3 用户输入与体验反馈", level=2)
    add_data_table(
        doc,
        ["用户输入", "输入方式", "系统如何使用", "用户看到的体验"],
        [
            ("城市", "北京 / 天津", "决定城市棋盘、角色库、路线和视觉主题", "进入《中轴入局》或《海河来信》。"),
            ("游玩时间", "60 / 90 / 150 分钟", "裁剪主线站点数量和支线开放数量", "短线不累，长线有隐藏支线。"),
            ("同行类型", "独自 / 情侣 / 朋友 / 亲子 / 研学", "调整任务难度、语气和协作要求", "亲子会多观察任务，朋友会多选择和合影任务。"),
            ("兴趣主题", "建筑 / 历史人物 / 市井生活 / 美食民俗 / 摄影", "决定角色追问和卡牌权重", "同一地点可以讲建筑，也可以讲生活。"),
            ("现场照片", "相机拍摄", "识别空间元素，不单独判断地点", "AI 会引用照片里的屋檐、桥、河面、街灯等细节。"),
            ("对话选择", "按钮或语音输入", "推进剧场分支并记录价值倾向", "角色会追问，最终故事会引用用户选择。"),
            ("一句留言", "文字或语音", "作为终局故事的个人记忆碎片", "生成游记中的“我留给这座城的话”。"),
        ],
        widths=[1200, 1450, 3300, 3170],
    )

    doc.add_heading("2.4 判定逻辑：地点不是只靠照片判断", level=2)
    add_callout(
        doc,
        "多信号触发",
        "地点确认采用 GPS/围栏为主，现场照片为剧情增强，天气/时间/用户选择为氛围调节。照片识别失败时，用户仍可通过“我已到达”按钮和手动选择建筑元素继续游戏。",
        fill=PALE_GREEN,
    )
    add_data_table(
        doc,
        ["信号", "用途", "失败时兜底"],
        [
            ("GPS/围栏", "确认用户在指定地点范围内", "给出步行引导；允许用户上传现场照片请求人工式确认。"),
            ("照片识别", "识别建筑、街景、天气、角度、颜色和人流", "让用户手动选择“我拍到了：屋檐/桥/河/街灯/牌楼”。"),
            ("天气/时间", "调整剧情语气和任务提示", "使用默认白天晴朗模板。"),
            ("用户选择", "决定角色追问、卡牌和终局主题", "提供 3 个预设选项，降低输入成本。"),
        ],
        widths=[1600, 4200, 3320],
    )

    add_section_intro(
        doc,
        "3. AI 剧场内容规则",
        "AI 剧场的目标不是自由聊天，而是在史料护栏内完成角色扮演。每个角色都由事实卡、人格卡、禁区卡组成。AI 可以生成表达、追问和情绪，但不能随意编造史实。",
    )
    doc.add_heading("3.1 三卡结构", level=2)
    add_data_table(
        doc,
        ["卡片", "内容", "作用"],
        [
            ("事实卡", "地点历史、建筑特征、时代背景、可核实事件", "保证角色发言有事实依据。"),
            ("人格卡", "身份、年龄感、说话方式、价值观、关心的问题", "让角色像一个“人”，不是百科。"),
            ("禁区卡", "不能编造、不能误导、不能浪漫化的问题", "控制幻觉和价值风险。"),
        ],
        widths=[1400, 4300, 3420],
    )
    doc.add_heading("3.2 对话长度", level=2)
    add_bullets(
        doc,
        [
            "每个地点默认 2-4 轮对话，总时长控制在 3-5 分钟。",
            "第一句话必须引用现场照片或当前位置，不用抽象开场。",
            "每一轮给用户 3 个明确选择，并允许自由输入。",
            "用户如果连续两轮不输入，系统自动转入任务并发放基础卡牌。",
            "角色不直接讲完整历史课，而是通过问题、处境和物件让用户参与。",
        ]
    )
    doc.add_heading("3.3 示例：角色卡模板", level=2)
    add_kv_table(
        doc,
        [
            ("角色名", "宫廷画师"),
            ("适用地点", "北京故宫角楼 / 筒子河外侧"),
            ("事实卡", "角楼、屋檐、水面、宫城空间、轴线秩序；避免声称用户进入未开放区域。"),
            ("人格卡", "克制、讲究构图、相信空间能塑造人的心；常用画面、线条、光影作比喻。"),
            ("禁区卡", "不编造具体宫廷秘闻；不把建筑细节说成确定的单一政治隐喻；不要求用户翻越护栏或进入危险区域。"),
            ("开场逻辑", "从用户照片中的屋檐、水面、天空、阴影或角楼轮廓切入。"),
        ],
        widths=[1600, 7520],
        fill=LIGHT_GRAY,
    )

    add_section_intro(
        doc,
        "4. 北京篇：《中轴入局》",
        "北京篇的主题是“秩序与烟火”。路线不依赖故宫入场预约，重点使用前门、故宫角楼外侧、景山、鼓楼/什刹海等公共可达区域。北京中轴线本身具有清晰的南北结构和都城秩序感，适合做成一盘可行走的文化棋局。",
    )
    doc.add_heading("4.1 北京主线概览", level=2)
    add_data_table(
        doc,
        ["顺序", "棋格", "地点", "角色", "卡牌", "建议停留"],
        [
            ("1", "起局：城门开市", "前门 / 正阳门外", "清末掌柜或城门吏", "城门卡 / 商贸卡", "12 分钟"),
            ("2", "中轴取景", "中轴视线点 / 天安门外侧视角", "营城匠师", "秩序卡", "10 分钟"),
            ("3", "宫城水影", "故宫角楼 / 筒子河外侧", "宫廷画师", "宫城卡", "15 分钟"),
            ("4", "登高观城", "景山", "观城史官", "俯瞰卡", "20 分钟"),
            ("5", "旧城回声", "鼓楼 / 什刹海", "更夫或胡同居民", "烟火卡", "15 分钟"),
        ],
        widths=[650, 1850, 2450, 1850, 1420, 900],
        font_size=8.5,
    )
    add_callout(
        doc,
        "北京体验关键词",
        "中轴、对称、城门、宫城、登高、胡同、生活。北京篇不是讲“皇城有多宏大”，而是让用户感受到一座城市如何通过空间安排人的行动，又如何在宏大秩序旁边长出日常烟火。",
        fill=PALE_YELLOW,
    )

    beijing_nodes = [
        {
            "title": "4.2.1 前门 / 正阳门外：城门开市",
            "place": "前门大街、正阳门外侧、可安全停留的步行区域。",
            "input": "用户选择“我看到的是城门/招牌/人流/老字号”；拍摄门楼、牌匾、街面或人群。",
            "ai": "GPS 确认前门范围；照片识别门楼、招牌、街道纵深、人流密度；时间决定早市/暮色氛围。",
            "role": "清末掌柜或城门吏。掌柜负责商贸和市井，城门吏负责城门与秩序。",
            "line": "“你拍到的不是一条街，是城门外最会说话的地方。人从这里进城，货从这里换手，消息也从这里散开。”",
            "task": "找一个能体现“买卖发生过”的细节：招牌、门脸、铺面、吆喝声、排队人群。拍照或录 10 秒声音。",
            "card": "商贸卡。若用户选择“秩序/城门”倾向，也可获得城门卡。",
            "fallback": "若照片识别不到门楼，系统让用户手动选择“我拍到的是：招牌/街道/门楼/人流”，继续触发剧场。",
        },
        {
            "title": "4.2.2 中轴视线点：中轴取景",
            "place": "天安门广场外围或可合法停留的中轴视线点，不要求进入管控区域。",
            "input": "用户拍摄一张尽量对称的照片，并选择“我想看秩序/仪式/城市尺度”。",
            "ai": "识别道路中线、建筑对称、广场开阔感、天空和人流；结合用户主题选择。",
            "role": "营城匠师。以尺度、方向、对称和规制说话。",
            "line": "“你把线拍直了。可真正难的不是画一条线，是让一座城愿意沿着这条线呼吸几百年。”",
            "task": "拍出一张“中轴感”照片：画面中有明确的正中线、左右平衡或远近层次。",
            "card": "秩序卡。",
            "fallback": "若现场不便拍摄，提供“观察题”：选择哪一个元素最能体现中轴秩序：道路、城楼、广场、视线。",
        },
        {
            "title": "4.2.3 故宫角楼 / 筒子河：宫城水影",
            "place": "故宫角楼、筒子河外侧、可拍摄角楼与水面的公共区域。",
            "input": "用户拍摄角楼、屋檐、水面或倒影；选择“我想问建筑/权力/美”。",
            "ai": "识别屋檐层次、水面倒影、天空颜色、视角远近；生成画师视角的开场。",
            "role": "宫廷画师。克制、审美敏感，常用构图、线条和光影回应用户。",
            "line": "“你拍到的不是一座楼，而是一种被安排好的秩序。水替它收住了倒影，屋檐替它收住了天空。”",
            "task": "拍一个“角楼与当下同框”的画面：古建筑 + 行人/树影/天空/城市生活。",
            "card": "宫城卡。",
            "fallback": "若人流太多或天气差，任务改为选择画面里的一个元素，让画师围绕它讲一段 80 字以内的独白。",
        },
        {
            "title": "4.2.4 景山：登高观城",
            "place": "景山公园内可登高观城区域；如当天无法进入，使用景山周边外侧替代剧情。",
            "input": "用户选择“俯瞰/风/城市层次/远处的线”；拍摄高处城市视野。",
            "ai": "识别视野开阔度、远近层次、天气和能见度；根据是否看到中轴线生成不同文本。",
            "role": "观城史官。既看空间，也看时间；语气沉稳。",
            "line": "“站高一点，城市就不只是街道了。你会看见它把宫城、街巷、树影和人的脚步，都排进同一张长卷。”",
            "task": "拍一张“城市层次”：近处树木或人，中景建筑，远处天空或轴线。",
            "card": "俯瞰卡。",
            "fallback": "若不能登高，系统改为“想象登高”剧场：用户选择一个方向，史官讲那个方向连接的城市记忆。",
        },
        {
            "title": "4.2.5 鼓楼 / 什刹海：旧城回声",
            "place": "鼓楼外侧、什刹海周边、胡同入口等可停留区域。",
            "input": "用户拍摄胡同、门牌、湖面、自行车、店铺或街声；留下一句话。",
            "ai": "识别生活元素和声音素材；将宏大中轴叙事转入普通人的日常。",
            "role": "更夫或胡同居民。更夫代表时间，胡同居民代表生活记忆。",
            "line": "“城楼上的钟声远，胡同里的锅铲近。你走到这里，棋局就不只属于帝王，也属于晚饭前回家的人。”",
            "task": "录制 10 秒城市声音，或拍一处“仍在使用的生活细节”。",
            "card": "烟火卡。完成留言后获得回声卡。",
            "fallback": "如果用户不想录音，可选择“我听见了：车铃/人声/水声/脚步/风声”，系统生成对应回声。",
        },
    ]
    for node in beijing_nodes:
        add_city_node(doc, node)

    doc.add_heading("4.3 北京终局生成", level=2)
    add_bullets(
        doc,
        [
            "若用户收集“秩序卡 + 宫城卡 + 俯瞰卡”，标题倾向为《一座城如何把时间排成一条线》。",
            "若用户收集“商贸卡 + 烟火卡 + 回声卡”，标题倾向为《我在北京看见秩序旁边的生活》。",
            "若用户自由留言偏情感，游记增加“写给旧城的一句话”章节。",
            "若用户摄影主题明显，海报突出照片构图与路线轨迹。",
        ]
    )

    add_section_intro(
        doc,
        "5. 天津篇：《海河来信》",
        "天津篇的主题是“流动与相遇”。它不复制北京的帝都叙事，而是围绕海河、桥、开埠、建筑、商贸和民俗展开。主线适合从天津站附近进入，沿海河与意式风情区、古文化街形成一条较紧凑的可步行体验；五大道作为支线，用于更长时间或建筑兴趣用户。",
    )
    doc.add_heading("5.1 天津主线概览", level=2)
    add_data_table(
        doc,
        ["顺序", "棋格", "地点", "角色", "卡牌", "建议停留"],
        [
            ("1", "起局：河岸抵达", "天津站 / 世纪钟 / 解放桥", "海河船工或报童", "开埠卡", "12 分钟"),
            ("2", "异乡街灯", "意式风情区 / 马可波罗广场", "建筑师或翻译青年", "洋楼卡", "15 分钟"),
            ("3", "海河传话", "海河岸边", "码头商人", "河岸卡", "10 分钟"),
            ("4", "津门故里", "古文化街 / 天后宫周边", "泥人张学徒或香客", "民俗卡", "18 分钟"),
            ("支线", "万国建筑", "五大道", "旧宅管家或外交青年", "万国建筑卡", "30-60 分钟"),
        ],
        widths=[650, 1750, 2450, 1850, 1420, 1000],
        font_size=8.5,
    )
    add_callout(
        doc,
        "天津体验关键词",
        "海河、桥、开埠、洋楼、码头、民俗、买卖、市声。天津篇不是把城市包装成“异国风情”，而是讲一个港口城市如何把远方、商业、语言和生活接到一起。",
        fill=PALE_YELLOW,
    )

    tianjin_nodes = [
        {
            "title": "5.2.1 天津站 / 世纪钟 / 解放桥：河岸抵达",
            "place": "天津站前广场、世纪钟、解放桥、海河可视区域。",
            "input": "用户拍摄桥、钟、车站、河面或人流；选择“我是刚抵达/重新回到/路过天津”。",
            "ai": "GPS 确认天津站-解放桥范围；照片识别桥梁、河面、钟表、车站动线。",
            "role": "海河船工或报童。船工讲河与货，报童讲消息和城市速度。",
            "line": "“你从车站出来，第一眼就看见河。天津总是这样，先让人抵达，再把远方交给水面。”",
            "task": "拍一张“抵达感”照片：桥、钟、车站或河面至少出现一个。",
            "card": "开埠卡。",
            "fallback": "若照片只拍到人流，报童角色启动，以“人潮和消息”进入剧情。",
        },
        {
            "title": "5.2.2 意式风情区 / 马可波罗广场：异乡街灯",
            "place": "意式风情区、马可波罗广场、街灯、拱窗、欧式立面周边。",
            "input": "用户拍摄拱窗、街灯、广场、餐厅外立面；选择“建筑/语言/远方/命运”。",
            "ai": "识别拱券、阳台、街灯、广场尺度、夜景或晴天色彩。",
            "role": "翻译青年或建筑师。翻译青年连接语言与命运，建筑师连接风格与城市更新。",
            "line": "“我每天在这条街上替不同语言的人传话。可有些话，翻译得出字，翻译不出命运。”",
            "task": "找到一个“不是本地传统街巷”的建筑细节：拱窗、栏杆、街灯、柱式或广场边界。",
            "card": "洋楼卡。",
            "fallback": "若用户无法识别建筑细节，系统给出四张元素按钮：拱窗、街灯、阳台、广场，由用户点选。",
        },
        {
            "title": "5.2.3 海河岸边：海河传话",
            "place": "意式风情区与古文化街之间的海河岸线，或附近安全亲水步道。",
            "input": "用户录制 10 秒河岸声音，或拍摄船、桥、倒影、栏杆。",
            "ai": "识别水面、桥、船、夜景灯光；声音任务可作为游记素材。",
            "role": "码头商人。关注货物、时间、交易、天气和消息。",
            "line": "“河面看着安静，其实每一道水纹都带过价钱、口音、货单和人的念头。”",
            "task": "录一段河岸声音，或拍一张“河把两岸连起来”的照片。",
            "card": "河岸卡。",
            "fallback": "若天气或安全不适合停留，改为远景拍摄任务，不要求靠近水边。",
        },
        {
            "title": "5.2.4 古文化街 / 天后宫周边：津门故里",
            "place": "古文化街牌楼、商铺、民俗手作、天后宫周边可开放区域。",
            "input": "用户拍摄牌楼、民俗纹样、泥塑、风筝、年画、店招；选择“手艺/祈愿/买卖/热闹”。",
            "ai": "识别仿古街区、民俗商品、纹样、牌楼和人流；结合用户选择生成民俗剧情。",
            "role": "泥人张学徒或香客。学徒讲手艺，香客讲祈愿和城市生活。",
            "line": "“别看这条街热闹，真正留下来的东西，常常在一双手里：捏泥、上色、写字、递香。”",
            "task": "拍一个“手作痕迹”：泥塑、笔墨、纸鸢、纹样、老字号包装或手工动作。",
            "card": "民俗卡。",
            "fallback": "若现场拥挤，用户可选择“我看见了哪种民俗元素”，系统生成短剧场并发放基础卡。",
        },
        {
            "title": "5.2.5 五大道支线：万国建筑",
            "place": "马场道、睦南道、大理道、常德道、重庆道及周边可步行区域。",
            "input": "用户选择建筑兴趣后开启；拍摄一栋不同风格的小楼、门窗、墙面或街树。",
            "ai": "识别洋楼、红砖、阳台、庭院、街道尺度；不强行判断具体建筑年代和人物归属。",
            "role": "旧宅管家或外交青年。讲近代生活、居住秩序和城市身份变化。",
            "line": "“这些房子看起来安静，其实每扇窗后都换过主人、语言、旗帜和晚餐的味道。”",
            "task": "拍三种建筑细节中的一种：窗、门、墙面。用户选择它像哪种情绪：克制、华丽、陌生、安稳。",
            "card": "万国建筑卡。",
            "fallback": "若时间不足，支线改为“建筑盲盒”：系统展示一段虚拟角色来信，不要求实地完成。",
        },
    ]
    for node in tianjin_nodes:
        add_city_node(doc, node)

    doc.add_heading("5.3 天津终局生成", level=2)
    add_bullets(
        doc,
        [
            "若用户收集“开埠卡 + 河岸卡 + 洋楼卡”，标题倾向为《天津把远方变成了街道》。",
            "若用户收集“民俗卡 + 河岸卡 + 回声卡”，标题倾向为《我在海河边收到一封旧时来信》。",
            "若用户完成五大道支线，游记加入“万国建筑支线章”。",
            "若用户录制声音素材，最终海报可生成“我的海河声纹”文案。",
        ]
    )

    add_section_intro(
        doc,
        "6. App 功能设计",
        "第一版功能围绕“可完成的一次城市剧情体验”设计，不做复杂社交、不做大型开放世界、不做重 AR。每个功能都服务于现场触发、剧情推进、卡牌收集和最终成果生成。",
    )
    doc.add_heading("6.1 页面结构", level=2)
    add_data_table(
        doc,
        ["页面", "用户看到什么", "用户能做什么"],
        [
            ("首页", "北京、天津两座城市入口；今日推荐路线", "选择城市，查看预计时长和主题。"),
            ("入局问答", "游玩时间、同行类型、兴趣主题、体力偏好", "输入偏好，生成路线。"),
            ("城市棋盘", "主线格子、支线格子、当前进度、下一站", "导航到下一站，查看已收集卡牌。"),
            ("地点触发页", "当前位置、拍照提示、时空骰按钮", "拍照、掷骰、进入剧场。"),
            ("AI 剧场页", "角色头像/声音、对话、选择按钮", "选择提问、语音输入、继续剧情。"),
            ("任务页", "现实任务说明、拍照/录音/选择控件", "完成任务，提交证据或选择。"),
            ("卡牌册", "已获得卡牌、卡牌组合、隐藏结局提示", "查看卡牌，理解最终故事方向。"),
            ("终局页", "个人游记、路线图、照片、角色、海报", "生成、编辑标题、分享。"),
        ],
        widths=[1300, 3900, 3920],
    )

    doc.add_heading("6.2 入局问答文案", level=2)
    add_kv_table(
        doc,
        [
            ("问题 1", "今天想进入哪座城市？选项：北京《中轴入局》 / 天津《海河来信》。"),
            ("问题 2", "你有多少时间？选项：60 分钟轻体验 / 90 分钟标准线 / 150 分钟完整线。"),
            ("问题 3", "这次和谁一起？选项：独自 / 情侣 / 朋友 / 亲子 / 研学小队。"),
            ("问题 4", "你更想追哪种回声？选项：建筑 / 历史人物 / 市井生活 / 美食民俗 / 摄影。"),
            ("问题 5", "今天的步行状态？选项：轻松走 / 正常走 / 可以多走一点。"),
        ],
        widths=[1600, 7520],
        fill=LIGHT_GRAY,
    )

    doc.add_heading("6.3 最终生成物", level=2)
    add_bullets(
        doc,
        [
            "《我的京津时空游记》：自动生成 800-1200 字，可编辑标题。",
            "城市记忆海报：包含城市、路线、卡牌、1-3 张用户照片和一句回声。",
            "文化卡牌册：展示本次获得卡牌和未解锁支线。",
            "路线回放：用时间线展示用户经过的地点、遇见的角色和做出的选择。",
        ]
    )

    add_section_intro(
        doc,
        "7. 技术与内容实现",
        "技术实现不追求一开始就做复杂 AR，而是先确保定位、拍照、AI 角色、任务和生成成果稳定。内容实现采用结构化剧本库，确保每个地点可以复用统一模板。",
    )
    doc.add_heading("7.1 系统模块", level=2)
    add_data_table(
        doc,
        ["模块", "职责", "第一版要求"],
        [
            ("路线引擎", "根据城市、时间、体力和开放情况生成路线", "支持固定主线、支线开关、站点跳过。"),
            ("地点触发器", "GPS/围栏确认用户到达地点", "半径范围可配置，提供手动兜底。"),
            ("照片理解", "识别建筑元素、天气、光线、角度和氛围", "输出标签，不单独决定地点。"),
            ("AI 剧场", "按角色卡和事实库生成对话", "每站 2-4 轮，引用现场元素。"),
            ("任务系统", "发布拍照、录音、观察、选择任务", "任务证据可轻量提交。"),
            ("卡牌系统", "根据完成情况发放文化卡牌", "卡牌影响终局故事主题。"),
            ("生成系统", "生成游记、海报、路线回放", "允许用户改标题和删除照片。"),
            ("内容 CMS", "管理城市、地点、角色、事实卡、禁区卡", "内容人员可配置，不依赖发版。"),
        ],
        widths=[1300, 3900, 3920],
    )

    doc.add_heading("7.2 地点配置模板", level=2)
    add_kv_table(
        doc,
        [
            ("location_id", "城市缩写 + 序号，例如 bj_03_corner_tower、tj_02_italian_area。"),
            ("geo_fence", "中心点坐标、半径、可替代触发点。"),
            ("photo_tags", "该地点推荐识别元素，如屋檐、水面、桥、拱窗、牌楼。"),
            ("role_card", "角色身份、说话方式、关心的问题。"),
            ("fact_card", "可引用事实点和资料来源。"),
            ("forbidden_card", "禁止编造、禁止误导、禁止引导危险行为。"),
            ("missions", "拍照、录音、观察、选择题任务列表。"),
            ("reward_cards", "完成基础任务和隐藏任务分别发放的卡牌。"),
            ("fallback", "定位失败、照片失败、天气不佳、场地关闭时的替代体验。"),
        ],
        widths=[1800, 7320],
        fill=LIGHT_GRAY,
    )

    doc.add_heading("7.3 隐私与安全", level=2)
    add_bullets(
        doc,
        [
            "用户照片默认只用于本次剧情生成和个人成果，需明确提示是否保存到云端。",
            "涉及人脸的照片可在上传前提示用户裁剪或模糊处理。",
            "路线提示必须避免穿越机动车道、翻越护栏、进入未开放区域。",
            "所有地点任务必须提供“不拍照也能继续”的替代选择。",
            "历史内容必须标注来源等级，AI 不确定时应使用“据资料显示/可能/传说中”这类谨慎表达。",
        ]
    )

    add_section_intro(
        doc,
        "8. 内容生产清单",
        "为完成首发 MVP，需要为北京和天津各生产一套最小内容包。不要一次性写几十个景点，先确保 10 个棋格全部可玩、可测、可讲清楚。",
    )
    add_data_table(
        doc,
        ["内容项", "北京数量", "天津数量", "说明"],
        [
            ("地点配置", "5 个主线", "4 个主线 + 1 个支线", "每个地点含围栏、拍照提示、任务和兜底。"),
            ("AI 角色", "5 个", "5 个", "每个角色含事实卡、人格卡、禁区卡。"),
            ("剧情事件", "每点 6 个骰面", "每点 6 个骰面", "同地点可重复游玩。"),
            ("文化卡牌", "6-8 张", "6-8 张", "含基础卡、隐藏卡、回声卡。"),
            ("任务", "每点 2-3 个", "每点 2-3 个", "拍照、录音、观察、选择混合。"),
            ("终局模板", "4 个", "4 个", "由卡牌组合和用户留言决定。"),
            ("分享海报模板", "2 个", "2 个", "城市专属视觉风格。"),
        ],
        widths=[1600, 1400, 1700, 4420],
    )

    doc.add_heading("8.1 卡牌体系", level=2)
    add_data_table(
        doc,
        ["城市", "卡牌", "含义", "影响结局"],
        [
            ("北京", "城门卡", "进入城市的边界与规则", "强化“入城/秩序”叙事。"),
            ("北京", "商贸卡", "前门市井与买卖", "强化“生活/交易”叙事。"),
            ("北京", "秩序卡", "中轴线、对称与都城理想", "强化“空间秩序”叙事。"),
            ("北京", "宫城卡", "宫城、审美与权力空间", "强化“建筑与制度”叙事。"),
            ("北京", "烟火卡", "胡同、湖面、日常声音", "强化“旧城生活”叙事。"),
            ("天津", "开埠卡", "抵达、港口、消息与远方", "强化“流动/抵达”叙事。"),
            ("天津", "洋楼卡", "异国建筑与城市更新", "强化“建筑/语言”叙事。"),
            ("天津", "河岸卡", "海河连接两岸与人群", "强化“河流/商业”叙事。"),
            ("天津", "民俗卡", "手艺、祈愿、商铺与传统", "强化“市井/民俗”叙事。"),
            ("天津", "万国建筑卡", "五大道支线与近代城市身份", "解锁建筑支线章节。"),
            ("通用", "回声卡", "用户留给城市的一句话", "决定游记结尾和海报主文案。"),
        ],
        widths=[1000, 1500, 3600, 3020],
        font_size=8.8,
    )

    add_section_intro(
        doc,
        "9. 试玩与路演脚本",
        "路演不要展示完整城市库，而要展示两个高辨识度场景：北京故宫角楼和天津意式风情区。每个场景 90 秒，让评委立刻理解照片、AI 剧场、任务、卡牌和最终生成物之间的关系。",
    )
    doc.add_heading("9.1 北京路演片段", level=2)
    add_numbers(
        doc,
        [
            "主持人选择北京《中轴入局》，输入 90 分钟、建筑、朋友同行。",
            "App 显示前门、故宫角楼、景山、鼓楼/什刹海路线。",
            "演示用户在故宫角楼外拍照。",
            "系统识别：角楼、屋檐、水面、阴天。",
            "时空骰掷出“风物”，宫廷画师登场。",
            "画师说：“你拍到的不是一座楼，而是一种被安排好的秩序。”",
            "用户选择“空间真的会改变人的心吗？”",
            "角色追问，用户完成“古建筑与当下同框”任务，获得宫城卡。",
            "终局预览生成《一座城如何把时间排成一条线》。",
        ]
    )
    doc.add_heading("9.2 天津路演片段", level=2)
    add_numbers(
        doc,
        [
            "主持人选择天津《海河来信》，输入 90 分钟、建筑、市井生活。",
            "App 显示天津站、解放桥、意式风情区、海河岸、古文化街路线。",
            "演示用户在意式风情区拍摄拱窗和街灯。",
            "系统识别：拱窗、街灯、欧式立面、广场。",
            "时空骰掷出“来信”，翻译青年登场。",
            "角色说：“有些话，翻译得出字，翻译不出命运。”",
            "用户选择“你最难翻译的一句话是什么？”",
            "完成建筑细节任务，获得洋楼卡。",
            "终局预览生成《天津把远方变成了街道》。",
        ]
    )

    add_section_intro(
        doc,
        "10. 六周执行计划",
        "六周计划以“能试走、能演示、能生成成果”为目标。每周都有可验收产物，避免陷入无限扩展城市和剧情。",
    )
    add_data_table(
        doc,
        ["周次", "重点", "交付物", "验收标准"],
        [
            ("第 1 周", "路线与内容定稿", "京津 10 个棋格配置表；角色三卡初稿", "每个地点都有触发、任务、卡牌和兜底。"),
            ("第 2 周", "App 原型", "首页、入局问答、城市棋盘、地点页、剧场页、卡牌册、终局页", "能完整点击走通一条假数据流程。"),
            ("第 3 周", "AI 剧场引擎", "照片标签输入、角色卡提示词、对话分支、时空骰", "同一地点能根据照片标签生成不同开场。"),
            ("第 4 周", "北京篇打通", "北京 5 点完整可玩", "从入局到游记生成无断点。"),
            ("第 5 周", "天津篇打通", "天津 4 主线 + 1 支线完整可玩", "支线可选，不影响主线完成。"),
            ("第 6 周", "实地测试与路演", "测试报告、Demo 视频、路演稿、分享海报样张", "10-20 人试玩，完成率、拍照成功率和分享意愿有数据。"),
        ],
        widths=[950, 1900, 3400, 2870],
        font_size=8.8,
    )

    doc.add_heading("10.1 关键指标", level=2)
    add_data_table(
        doc,
        ["指标", "目标值", "说明"],
        [
            ("主线完成率", ">= 70%", "用户至少完成 4 个主线棋格。"),
            ("拍照成功率", ">= 85%", "照片识别或手动兜底能成功进入剧场。"),
            ("剧场互动率", ">= 60%", "用户至少主动选择或输入 3 次。"),
            ("终局生成率", ">= 80%", "完成后成功生成游记或海报。"),
            ("分享意愿", ">= 30%", "愿意保存或分享海报。"),
            ("内容可信度反馈", ">= 4/5", "用户认为内容不像胡编。"),
        ],
        widths=[2000, 1500, 5620],
    )

    add_section_intro(
        doc,
        "11. 风险与兜底",
        "实景文旅游戏最大的风险不是创意不足，而是现场不可控。第一版必须把天气、定位、照片识别、场地关闭、用户疲劳和历史幻觉都提前设计好。",
    )
    add_data_table(
        doc,
        ["风险", "表现", "兜底方案"],
        [
            ("天气不佳", "用户不愿停留或拍照效果差", "自动切换“雨天/夜色/风声”剧情；减少拍照要求。"),
            ("场地管控", "某点临时不可达", "路线引擎跳过该点，发放“错过的来信”虚拟剧情。"),
            ("照片识别失败", "AI 看不出建筑元素", "用户手动选择元素，仍能进入剧场。"),
            ("用户不想说话", "对话中断", "提供按钮选择；两轮无输入后自动进入任务。"),
            ("历史幻觉", "AI 编造事实", "角色卡 + 事实库 + 禁区卡；不确定内容用谨慎表达。"),
            ("路线太累", "用户中途退出", "提供 60 分钟短线和“一键终局”生成未完成版游记。"),
            ("隐私担忧", "用户不愿上传照片", "本地预览、可删除照片、可用文字选择替代照片。"),
        ],
        widths=[1600, 2900, 4620],
    )

    add_section_intro(
        doc,
        "12. 资料依据与内容边界",
        "本文档中的地点与主题以公开资料和可实地体验的旅游空间为基础。正式上线前，所有历史事实、开放时间、门票和路线安全信息都需要由内容负责人二次核验，并在 CMS 中标注来源。",
    )
    add_bullets(
        doc,
        [
            "北京中轴线：作为北京历史城区的南北轴线，串联宫殿、坛庙、城门、道路遗存等，是北京篇的核心叙事母题。",
            "天津古文化街：作为天津民俗、旅游、商业和手作文化集中体验空间，适合作为天津篇的民俗终局。",
            "天津意式风情区：以街区、广场、拱窗、街灯和近代建筑风貌承接“语言、远方与城市相遇”的剧场主题。",
            "五大道：作为天津近代建筑和街区风貌的代表区域，适合做建筑兴趣用户的可选支线。",
            "所有角色均为文化叙事角色或复合型虚构角色，不冒充具体历史人物的真实言行。",
        ]
    )
    p = doc.add_paragraph()
    set_para_format(p, after=3, line=1.2)
    set_run_font(p.add_run("参考链接："), 10.5, True, DARK_BLUE)
    links = [
        ("UNESCO 北京中轴线世界遗产页面", "https://whc.unesco.org/en/list/1714"),
        ("北京市政府英文网：北京中轴线", "https://english.beijing.gov.cn/specials/centralaxis/thepastandpresentcentralaxis/?v=jt"),
        ("China Daily：天津古文化街旅游区", "https://govt.chinadaily.com.cn/s/201905/17/WS5cde10ed498e079e68021151/tianjin-ancient-culture-street-tourist-area.html"),
        ("Visit Beijing：天津意式风情街", "https://english.visitbeijing.com.cn/article/47OO1rzB2hc"),
    ]
    for text, url in links:
        p = doc.add_paragraph(style="List Bullet")
        set_para_format(p, after=2, line=1.15)
        add_hyperlink(p, text, url)

    add_callout(
        doc,
        "最终执行原则",
        "先让一条北京线和一条天津线真的好玩，再谈更多城市。产品扩展的核心不是堆景点，而是复制“地点触发—角色剧场—现实任务—卡牌—个人故事”的内容生产能力。",
        fill=PALE_GREEN,
        spacer_after=False,
    )

    doc.save(OUT_PATH)
    return OUT_PATH


if __name__ == "__main__":
    path = build_document()
    print(path)
